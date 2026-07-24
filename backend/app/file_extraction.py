import re
from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader


ALLOWED_EXTENSIONS = {".pdf": "pdf", ".docx": "docx", ".txt": "txt"}
ALLOWED_MIME_TYPES = {
    ".pdf": {"application/pdf", "application/octet-stream"},
    ".docx": {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/zip",
        "application/octet-stream",
    },
    ".txt": {"text/plain", "application/octet-stream"},
}
MIN_EXTRACTED_CHARACTERS = 20


class FileExtractionError(ValueError):
    pass


def safe_filename(filename: str | None) -> str:
    name = Path((filename or "upload").replace("\\", "/")).name
    cleaned = re.sub(r"[\x00-\x1f<>:\"/\\|?*]+", "_", name).strip(" .")
    return cleaned[:180] or "upload"


def _normalise_text(text: str) -> str:
    lines = []
    for line in text.replace("\r\n", "\n").replace("\r", "\n").splitlines():
        cleaned = re.sub(r"[ \t]+", " ", line).strip()
        if cleaned:
            lines.append(cleaned)
        elif lines and lines[-1] != "":
            lines.append("")
    return "\n".join(lines).strip()


def _extract_pdf(content: bytes) -> str:
    if not content.startswith(b"%PDF"):
        raise FileExtractionError("The uploaded file is not a valid PDF.")
    try:
        reader = PdfReader(BytesIO(content))
        text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as error:
        raise FileExtractionError("Relay could not read this PDF.") from error
    if len(_normalise_text(text)) < MIN_EXTRACTED_CHARACTERS:
        raise FileExtractionError(
            "This PDF does not contain readable text. Paste the text manually or upload a text-based PDF."
        )
    return text


def _extract_docx(content: bytes) -> str:
    if not content.startswith(b"PK"):
        raise FileExtractionError("The uploaded file is not a valid DOCX document.")
    try:
        document = Document(BytesIO(content))
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n\n".join(parts)
    except Exception as error:
        raise FileExtractionError("Relay could not read this DOCX document.") from error


def _extract_txt(content: bytes) -> str:
    try:
        return content.decode("utf-8-sig")
    except UnicodeDecodeError as error:
        raise FileExtractionError("TXT files must use UTF-8 text encoding.") from error


def extract_uploaded_text(
    *,
    filename: str | None,
    content_type: str | None,
    content: bytes,
    max_bytes: int,
) -> tuple[str, str, str]:
    display_name = safe_filename(filename)
    extension = Path(display_name).suffix.casefold()
    if extension not in ALLOWED_EXTENSIONS:
        raise FileExtractionError("Supported file types are PDF, DOCX, and TXT.")
    if not content:
        raise FileExtractionError("The uploaded file is empty.")
    if len(content) > max_bytes:
        raise FileExtractionError(f"The uploaded file exceeds the {max_bytes // 1_048_576} MB limit.")
    mime = (content_type or "application/octet-stream").split(";")[0].strip().casefold()
    if mime not in ALLOWED_MIME_TYPES[extension]:
        raise FileExtractionError("The file content type does not match a supported document format.")

    extractors = {".pdf": _extract_pdf, ".docx": _extract_docx, ".txt": _extract_txt}
    text = _normalise_text(extractors[extension](content))
    if len(text) < MIN_EXTRACTED_CHARACTERS:
        raise FileExtractionError("We could not read usable text from this file. Paste the text manually instead.")
    return display_name, ALLOWED_EXTENSIONS[extension], text
